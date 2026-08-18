"""Genera el guion del chatbot en formato Word (.docx), ordenado y con
negrita en todo lo que la persona debe rellenar (lo que antes iba entre <>)."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

AZUL = RGBColor(0x00, 0x53, 0xE2)   # azul Walmart, para titulos
GRIS = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# --- estilo base ---
base = doc.styles["Normal"]
base.font.name = "Calibri"
base.font.size = Pt(11)


def titulo(texto, nivel=1):
    h = doc.add_heading(texto, level=nivel)
    for run in h.runs:
        run.font.color.rgb = AZUL
    return h


def parrafo(texto, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return p


def campo(etiqueta, ejemplo):
    """Una linea: 'Etiqueta:' normal + el texto a rellenar en NEGRITA."""
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(etiqueta + " ")
    r1.bold = False
    r2 = p.add_run(ejemplo)
    r2.bold = True
    r2.font.color.rgb = AZUL
    return p


# =========================================================================
# PORTADA / INTRO
# =========================================================================
t = doc.add_heading("Guion del Chatbot de WhatsApp", level=0)
for run in t.runs:
    run.font.color.rgb = AZUL

parrafo("Fase 1: menu, preguntas frecuentes (FAQ) y agendar citas.",
        italic=True, color=GRIS)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Como usar este documento: ").bold = True
p.add_run("todo lo que aparece ")
r = p.add_run("EN NEGRITA AZUL")
r.bold = True
r.font.color.rgb = AZUL
p.add_run(" es lo que TU debes cambiar por tu propio texto. "
          "El resto es la estructura, dejala tal cual. Cuando termines, "
          "me devuelves el documento y yo lo conecto al chatbot.")

p = doc.add_paragraph()
p.add_run("Variables magicas: ").bold = True
p.add_run("puedes escribir ")
for v in ["{nombre}", "{servicio}", "{dia}", "{hora}"]:
    rr = p.add_run(v + "  ")
    rr.bold = True
    rr.font.color.rgb = RGBColor(0x2A, 0x8A, 0x3E)  # verde
p.add_run("dentro de cualquier mensaje. El bot las reemplaza solo por los "
          "datos reales de la persona.")

# =========================================================================
# CONFIGURACION
# =========================================================================
titulo("1. Configuracion general", 1)
campo("Nombre del bot:", "Ej: Asistente de Peluqueria Luna")
campo("Tono de la conversacion:", "Ej: cercano / formal / juvenil")

# =========================================================================
# MENU PRINCIPAL
# =========================================================================
titulo("2. Menu principal (lo primero que ve la persona)", 1)
campo("Saludo + que puede hacer el bot:",
      "Ej: Hola! Soy {nombre_bot}, te ayudo con dudas y a agendar tu hora.")
parrafo("Texto de cada opcion del menu:", italic=True, color=GRIS)
campo("Opcion 1:", "Ej: Preguntas frecuentes")
campo("Opcion 2:", "Ej: Agendar una cita")
campo("Opcion 3:", "Ej: Informacion de contacto")
campo("Opcion 4:", "Ej: Hablar con una persona")
campo("Opcion 0:", "Ej: Salir")

# =========================================================================
# FAQ
# =========================================================================
titulo("3. Preguntas frecuentes (FAQ)", 1)
campo("Intro del menu de FAQ:", "Ej: Claro! Estas son las dudas mas comunes:")
parrafo("Titulos que aparecen en el submenu:", italic=True, color=GRIS)
campo("Titulo pregunta 1:", "Ej: Cuales son los horarios?")
campo("Titulo pregunta 2:", "Ej: Donde estan ubicados?")
campo("Titulo pregunta 3:", "Ej: Que servicios ofrecen?")
campo("Titulo pregunta 4:", "Ej: Cuanto cuestan?")
parrafo("Respuestas de cada pregunta:", italic=True, color=GRIS)
campo("Respuesta HORARIOS:", "Ej: Atendemos de lunes a viernes de 9 a 18h.")
campo("Respuesta UBICACION:", "Ej: Estamos en Av. Siempreviva 123, Santiago.")
campo("Respuesta SERVICIOS:", "Ej: Ofrecemos corte, color y peinado.")
campo("Respuesta PRECIOS:", "Ej: Corte desde $10.000, color desde $25.000.")

# =========================================================================
# INFO / CONTACTO
# =========================================================================
titulo("4. Informacion / contacto", 1)
campo("Datos de contacto:",
      "Ej: Telefono +56 9 1234 5678, correo hola@luna.cl, IG @luna")

# =========================================================================
# AGENDAR CITA
# =========================================================================
titulo("5. Agendar cita (flujo paso a paso)", 1)
parrafo("Paso 1 - Preguntar el servicio:", italic=True, color=GRIS)
campo("Mensaje:", "Ej: Perfecto! Que servicio quieres agendar?")
campo("Servicio opcion 1:", "Ej: Corte de pelo")
campo("Servicio opcion 2:", "Ej: Color")
campo("Servicio opcion 3:", "Ej: Peinado")
parrafo("Paso 2 - Preguntar el dia:", italic=True, color=GRIS)
campo("Mensaje:", "Ej: Genial! Para que dia te gustaria?")
parrafo("Paso 3 - Preguntar la hora:", italic=True, color=GRIS)
campo("Mensaje:", "Ej: A que hora te acomoda?")
parrafo("Paso 4 - Preguntar el nombre:", italic=True, color=GRIS)
campo("Mensaje:", "Ej: A nombre de quien reservamos?")
parrafo("Paso 5 - Confirmacion (usa las variables):", italic=True, color=GRIS)
campo("Mensaje:",
      "Ej: Perfecto {nombre}! Confirmo: {servicio} el {dia} a las {hora}. Correcto?")
parrafo("Mensaje final cuando la cita queda lista:", italic=True, color=GRIS)
campo("Mensaje:",
      "Ej: Listo {nombre}! Tu cita quedo agendada el {dia} a las {hora}. Te esperamos!")

# =========================================================================
# MENSAJES ESPECIALES
# =========================================================================
titulo("6. Mensajes especiales", 1)
campo("Derivar a una persona (humano):",
      "Ej: Claro, te derivo con alguien del equipo. En breve te contactan.")
campo("Despedida:", "Ej: Gracias por escribir! Que tengas un excelente dia.")
campo("Cuando NO entiende lo escrito:",
      "Ej: Ups, no te entendi bien. Puedes elegir una opcion del menu?")

doc.add_paragraph()
parrafo("Fin. Cambia todo lo que este en negrita azul y devuelveme el documento.",
        italic=True, color=GRIS)

salida = "chatbot/flujo/guion_planilla.docx"
doc.save(salida)
print("Documento generado:", salida)
